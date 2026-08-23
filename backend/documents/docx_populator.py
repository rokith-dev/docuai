from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from backend.documents.format_config import DocumentFormatConfig
from backend.documents.font_manager import FontManager


class DocxPopulator:
    """
    Populate DOCX templates while preserving the
    template's formatting and respecting font
    requirements detected from the template.
    """

    def __init__(
        self,
        format_config: DocumentFormatConfig | None = None,
    ):
        self.format_config = (
            format_config
            or DocumentFormatConfig()
        )

        self.default_font = FontManager.resolve(
            self.format_config.font_name
        )

    # ==================================================
    # MAIN
    # ==================================================

    def populate(
        self,
        template_path: str,
        output_path: str,
        template_map: dict,
        content: dict,
    ) -> str:

        template = Path(template_path)
        output = Path(output_path)

        if not template.exists():
            raise FileNotFoundError(
                f"Template file not found: {template_path}"
            )

        if template.suffix.lower() != ".docx":
            raise ValueError(
                "Only .docx templates are supported."
            )

        document = Document(
            str(template)
        )

        # ----------------------------------------------
        # Detect the document font
        # ----------------------------------------------

        detected_font = (
            self._detect_document_font(
                document
            )
        )

        self.document_font = (
            detected_font
            or self.default_font
        )

        # ----------------------------------------------
        # Populate paragraph fields
        # ----------------------------------------------

        self._populate_paragraph_fields(
            document,
            template_map,
            content,
        )

        # ----------------------------------------------
        # Populate table fields
        # ----------------------------------------------

        self._populate_table_fields(
            document,
            template_map,
            content,
        )

        # ----------------------------------------------
        # Save
        # ----------------------------------------------

        output.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document.save(
            str(output)
        )

        return str(output)

    # ==================================================
    # FONT DETECTION
    # ==================================================

    def _detect_document_font(
        self,
        document,
    ) -> str | None:
        """
        Detect the main font requested/used by the
        template.

        Example:
            Times New Roman
            Arial
            Calibri

        The first meaningful font used by the
        template is treated as the document font.
        """

        detected_fonts = []

        # ------------------------------------------
        # Paragraph runs
        # ------------------------------------------

        for paragraph in document.paragraphs:

            for run in paragraph.runs:

                font_name = (
                    run.font.name
                )

                if font_name:
                    detected_fonts.append(
                        font_name
                    )

        # ------------------------------------------
        # Table runs
        # ------------------------------------------

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        for run in paragraph.runs:

                            font_name = (
                                run.font.name
                            )

                            if font_name:
                                detected_fonts.append(
                                    font_name
                                )

        # ------------------------------------------
        # Look for explicit font names in text
        # ------------------------------------------

        all_text = []

        for paragraph in document.paragraphs:
            all_text.append(
                paragraph.text
            )

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    all_text.append(
                        cell.text
                    )

        combined_text = " ".join(
            all_text
        ).lower()

        known_fonts = [
            "times new roman",
            "arial",
            "calibri",
            "cambria",
            "georgia",
            "verdana",
            "tahoma",
            "courier new",
        ]

        for font in known_fonts:

            if font in combined_text:

                return font.title()

        # ------------------------------------------
        # Most common actual template font
        # ------------------------------------------

        if detected_fonts:

            counts = {}

            for font in detected_fonts:

                normalized = (
                    str(font).strip()
                )

                counts[normalized] = (
                    counts.get(
                        normalized,
                        0,
                    )
                    + 1
                )

            return max(
                counts,
                key=counts.get,
            )

        return None

    # ==================================================
    # PARAGRAPH FIELDS
    # ==================================================

    def _populate_paragraph_fields(
        self,
        document,
        template_map,
        content,
    ):

        for field in template_map.get(
            "fields",
            [],
        ):

            field_name = field[
                "name"
            ]

            if field_name not in content:
                continue

            location = field.get(
                "location",
                {},
            )

            if location.get(
                "source"
            ) != "paragraph":
                continue

            paragraph_index = location.get(
                "content_index"
            )

            if paragraph_index is None:
                continue

            if paragraph_index >= len(
                document.paragraphs
            ):
                continue

            paragraph = document.paragraphs[
                paragraph_index
            ]

            value = content[
                field_name
            ]

            if isinstance(
                value,
                dict,
            ):

                value = value.get(
                    "content",
                    "",
                )

            value = str(
                value
            )

            if not value.strip():
                continue

            content_type = field.get(
                "content_type",
                "text",
            )

            if content_type == "image":

                self._replace_paragraph_with_image(
                    paragraph,
                    value,
                )

            else:

                self._replace_paragraph_content(
                    paragraph,
                    value,
                    content_type,
                )

    # ==================================================
    # PARAGRAPH CONTENT
    # ==================================================

    def _replace_paragraph_content(
        self,
        paragraph,
        value,
        content_type,
    ):

        style = self._capture_run_style(
            paragraph
        )

        if content_type == "code":

            value = str(
                value
            )

        else:

            value = self._clean_text(
                value
            )

        self._remove_runs(
            paragraph
        )

        if content_type != "code":

            value = self._apply_capitalization(
                value,
                style.get(
                    "capitalization"
                ),
            )

        run = paragraph.add_run(
            value
        )

        self._apply_run_style(
            run,
            style,
        )

        if content_type == "code":

            run.font.size = Pt(
                self.format_config.code_font_size
            )

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )

        else:

            if not style.get(
                "size"
            ):

                run.font.size = Pt(
                    self.format_config.body_font_size
                )

            paragraph.alignment = (
                self._get_alignment()
            )

    # ==================================================
    # TABLE FIELDS
    # ==================================================

    def _populate_table_fields(
        self,
        document,
        template_map,
        content,
    ):

        fields = template_map.get(
            "table_fields",
            [],
        )

        used_locations = set()

        for field in fields:

            field_name = field[
                "name"
            ]

            if field_name not in content:
                continue

            value = content[
                field_name
            ]

            if isinstance(
                value,
                dict,
            ):

                value = value.get(
                    "content",
                    "",
                )

            value = str(
                value
            )

            if not value.strip():
                continue

            location = field.get(
                "value_location"
            )

            if not location:
                continue

            # --------------------------------------
            # Prevent two different fields from
            # overwriting the same table cell.
            # --------------------------------------

            location_key = (
                location.get(
                    "table_index"
                ),
                location.get(
                    "row"
                ),
                location.get(
                    "column"
                ),
            )

            if location_key in used_locations:

                location = (
                    self._find_alternative_table_location(
                        document,
                        field,
                        used_locations,
                    )
                )

                if location is None:
                    continue

            used_locations.add(
                (
                    location.get(
                        "table_index"
                    ),
                    location.get(
                        "row"
                    ),
                    location.get(
                        "column"
                    ),
                )
            )

            self._set_table_cell_value(
                document,
                location,
                value,
                field_name,
            )

    # ==================================================
    # FIND ALTERNATIVE TABLE LOCATION
    # ==================================================

    def _find_alternative_table_location(
        self,
        document,
        field,
        used_locations,
    ):
        """
        Try to find a nearby empty table cell
        when the detector accidentally maps two
        fields to the same cell.

        This is especially useful for Word tables
        containing merged cells.
        """

        label = str(
            field.get(
                "label",
                ""
            )
        ).strip().lower()

        location = field.get(
            "value_location"
        )

        if not location:
            return None

        table_index = location.get(
            "table_index"
        )

        row_index = location.get(
            "row"
        )

        column_index = location.get(
            "column"
        )

        if table_index >= len(
            document.tables
        ):
            return None

        table = document.tables[
            table_index
        ]

        # ------------------------------------------
        # Search nearby cells
        # ------------------------------------------

        candidates = []

        for r in range(
            max(
                0,
                row_index - 1,
            ),
            min(
                len(table.rows),
                row_index + 2,
            ),
        ):

            for c in range(
                max(
                    0,
                    column_index - 1,
                ),
                len(
                    table.rows[r].cells
                ),
            ):

                key = (
                    table_index,
                    r,
                    c,
                )

                if key in used_locations:
                    continue

                cell_text = (
                    table.rows[r]
                    .cells[c]
                    .text
                    .strip()
                    .lower()
                )

                # Empty cell is preferred.
                if not cell_text:

                    candidates.append(
                        (
                            r,
                            c,
                            0,
                        )
                    )

                # Don't overwrite a cell containing
                # another label.
                elif label not in cell_text:

                    candidates.append(
                        (
                            r,
                            c,
                            1,
                        )
                    )

        if not candidates:
            return None

        candidates.sort(
            key=lambda item: (
                item[2],
                abs(
                    item[0]
                    - row_index
                )
                + abs(
                    item[1]
                    - column_index
                ),
            )
        )

        r, c, _ = candidates[0]

        return {
            "table_index": table_index,
            "row": r,
            "column": c,
        }

    # ==================================================
    # SET TABLE CELL
    # ==================================================

    def _set_table_cell_value(
        self,
        document,
        location,
        value,
        field_name,
    ):

        table_index = location[
            "table_index"
        ]

        row_index = location[
            "row"
        ]

        column_index = location[
            "column"
        ]

        if table_index >= len(
            document.tables
        ):
            return

        table = document.tables[
            table_index
        ]

        if row_index >= len(
            table.rows
        ):
            return

        row = table.rows[
            row_index
        ]

        if column_index >= len(
            row.cells
        ):
            return

        cell = row.cells[
            column_index
        ]

        self._replace_cell_text(
            cell,
            value,
            field_name,
        )

    # ==================================================
    # TABLE CELL FORMATTING
    # ==================================================

    def _replace_cell_text(
        self,
        cell,
        value,
        field_name,
    ):

        if not cell.paragraphs:

            cell.text = (
                self._clean_text(
                    value
                )
            )

            return

        paragraph = cell.paragraphs[
            0
        ]

        style = self._capture_run_style(
            paragraph
        )

        original_alignment = (
            paragraph.alignment
        )

        value = self._clean_text(
            value
        )

        value = self._apply_capitalization(
            value,
            style.get(
                "capitalization"
            ),
        )

        self._remove_runs(
            paragraph
        )

        run = paragraph.add_run(
            value
        )

        self._apply_run_style(
            run,
            style,
        )

        # ------------------------------------------
        # Title fields are centered
        # ------------------------------------------

        if self._is_title_field(
            field_name
        ):

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

        else:

            paragraph.alignment = (
                original_alignment
            )

    # ==================================================
    # TITLE DETECTION
    # ==================================================

    @staticmethod
    def _is_title_field(
        field_name: str,
    ) -> bool:

        normalized_name = str(
            field_name
        ).strip().lower()

        return (
            normalized_name == "title"
            or "title" in normalized_name
        )

    # ==================================================
    # CLEAN TEXT
    # ==================================================

    @staticmethod
    def _clean_text(
        value: str,
    ) -> str:

        value = str(
            value
        )

        value = value.replace(
            "\t",
            " ",
        )

        value = " ".join(
            value.split()
        )

        return value.strip()

    # ==================================================
    # CAPTURE STYLE
    # ==================================================

    def _capture_run_style(
        self,
        paragraph,
    ):

        if not paragraph.runs:

            return {
                "font_name": self.document_font,
                "size": None,
                "bold": False,
                "italic": False,
                "underline": False,
                "color": None,
                "capitalization": "normal",
            }

        source_run = None

        for run in paragraph.runs:

            if run.text.strip():

                source_run = run
                break

        if source_run is None:

            source_run = paragraph.runs[
                0
            ]

        color = None

        try:

            if (
                source_run.font.color.type
                is not None
                and source_run.font.color.rgb
            ):

                color = str(
                    source_run.font.color.rgb
                )

        except Exception:

            color = None

        # ------------------------------------------
        # IMPORTANT:
        # Use detected document font if the
        # template specifies one.
        # ------------------------------------------

        font_name = (
            source_run.font.name
            or self.document_font
            or self.default_font
        )

        return {
            "font_name": font_name,

            "size": (
                source_run.font.size.pt
                if source_run.font.size
                else None
            ),

            "bold": (
                source_run.bold
                if source_run.bold is not None
                else False
            ),

            # Generated content should NOT inherit
            # italic instruction formatting.
            "italic": False,

            "underline": False,

            "color": color,

            "capitalization": (
                self._detect_capitalization(
                    source_run.text
                )
            ),
        }

    # ==================================================
    # APPLY STYLE
    # ==================================================

    def _apply_run_style(
        self,
        run,
        style,
    ):

        font_name = (
            style.get(
                "font_name"
            )
            or self.document_font
            or self.default_font
        )

        run.font.name = font_name

        if style.get(
            "size"
        ):

            run.font.size = Pt(
                style["size"]
            )

        else:

            run.font.size = Pt(
                self.format_config.body_font_size
            )

        run.bold = style.get(
            "bold",
            False,
        )

        # ------------------------------------------
        # Never copy italic instruction formatting
        # ------------------------------------------

        run.italic = False

        run.underline = False

        color = style.get(
            "color"
        )

        if color:

            try:

                run.font.color.rgb = (
                    RGBColor.from_string(
                        color
                    )
                )

            except ValueError:

                pass

    # ==================================================
    # CAPITALIZATION
    # ==================================================

    @staticmethod
    def _detect_capitalization(
        text: str,
    ) -> str:

        text = text.strip()

        if not text:
            return "normal"

        letters = [
            character
            for character in text
            if character.isalpha()
        ]

        if not letters:
            return "normal"

        if all(
            character.isupper()
            for character in letters
        ):

            return "uppercase"

        if all(
            character.islower()
            for character in letters
        ):

            return "lowercase"

        words = text.split()

        if words and all(
            word[:1].isupper()
            for word in words
            if word
        ):

            return "titlecase"

        return "normal"

    # ==================================================
    # APPLY CAPITALIZATION
    # ==================================================

    @staticmethod
    def _apply_capitalization(
        value: str,
        capitalization: str | None,
    ) -> str:

        if capitalization == "uppercase":

            return value.upper()

        if capitalization == "lowercase":

            return value.lower()

        if capitalization == "titlecase":

            return value.title()

        return value

    # ==================================================
    # REMOVE RUNS
    # ==================================================

    @staticmethod
    def _remove_runs(
        paragraph,
    ):

        for run in list(
            paragraph.runs
        ):

            run_element = run._element

            run_element.getparent().remove(
                run_element
            )

    # ==================================================
    # IMAGE
    # ==================================================

    def _replace_paragraph_with_image(
        self,
        paragraph,
        image_path,
    ):

        image = Path(
            image_path
        )

        if not image.exists():

            raise FileNotFoundError(
                f"Image file not found: {image_path}"
            )

        self._remove_runs(
            paragraph
        )

        run = paragraph.add_run()

        run.add_picture(
            str(image),
            width=Inches(5.5),
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    # ==================================================
    # BODY ALIGNMENT
    # ==================================================

    def _get_alignment(
        self,
    ):

        alignment = (
            self.format_config.body_alignment
            .upper()
            .strip()
        )

        if alignment == "CENTER":

            return WD_ALIGN_PARAGRAPH.CENTER

        if alignment == "RIGHT":

            return WD_ALIGN_PARAGRAPH.RIGHT

        if alignment == "LEFT":

            return WD_ALIGN_PARAGRAPH.LEFT

        return WD_ALIGN_PARAGRAPH.JUSTIFY