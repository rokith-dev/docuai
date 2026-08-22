from pathlib import Path

from docx import Document


class TemplateAnalyzer:
    """Analyze the structure of a DOCX template."""

    def analyze(self, file_path: str) -> dict:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Template file not found: {file_path}"
            )

        if path.suffix.lower() != ".docx":
            raise ValueError("Only .docx templates are supported.")

        document = Document(file_path)

        paragraphs = []

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()

            if not text:
                continue

            paragraphs.append(
                {
                    "index": index,
                    "text": text,
                    "style": paragraph.style.name,
                    "alignment": (
                        paragraph.alignment.name
                        if paragraph.alignment
                        else None
                    ),
                }
            )

        tables = []

        for table_index, table in enumerate(document.tables):
            rows = []

            for row in table.rows:
                rows.append(
                    [cell.text.strip() for cell in row.cells]
                )

            tables.append(
                {
                    "index": table_index,
                    "rows": rows,
                }
            )

        return {
            "file_name": path.name,
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "table_count": len(tables),
            "tables": tables,
        }